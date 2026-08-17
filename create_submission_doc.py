from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION

OUT = "submission/Madhavik_Dogra_Founding_AI_Engineer_Assignment.docx"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(29, 37, 48)
MUTED = RGBColor(94, 104, 117)
PALE = "E8EEF5"

def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd"); tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width)); tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)

def add_para(doc, text="", *, size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align: p.alignment = align
    r = p.add_run(text); set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p

doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11); normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1
for name, size, color, before, after in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK,8,4)]:
    s = styles[name]; s.font.name="Calibri"; s._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); s._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri"); s.font.size=Pt(size); s.font.color.rgb=color; s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

header = section.header.paragraphs[0]
header.text = "FOUNDING AI ENGINEER ASSIGNMENT                                      CHANGE SCOPE"
header.paragraph_format.space_after = Pt(0)
for run in header.runs: set_font(run, size=8.5, color=MUTED, bold=True)
add_page_number(section.footer.paragraphs[0])

# First page / memo masthead.
add_para(doc, "FOUNDING AI ENGINEER ASSIGNMENT", size=10, color=BLUE, bold=True, after=9)
title = add_para(doc, "ChangeScope", size=30, color=INK, bold=True, after=3)
title.paragraph_format.space_before = Pt(14)
add_para(doc, "An evidence-first planner for safe, multi-file engineering changes", size=14, color=MUTED, after=18)
for label, value in [("Candidate", "Madhavik Dogra"),("Repository", "https://github.com/madhav1k/changescope"),("Deployment", "https://changescope.vercel.app")]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    set_font(p.add_run(f"{label}: "), size=10.5, color=DARK, bold=True)
    set_font(p.add_run(value), size=10.5, color=INK)

heading(doc, "What I built and why", 1)
add_para(doc, "I built ChangeScope, a focused interface for turning a product request into a safe, reviewable implementation plan. It deliberately starts before code generation: a developer selects an issue, sees the likely dependency path, understands which files are relevant, reviews the plan, and approves it before execution.")
add_para(doc, "I chose this because the difficult part of agentic engineering is not producing a patch. It is deciding what deserves context, showing why a recommendation is trustworthy, and making the final action controllable. The prototype makes those steps visible instead of hiding them behind a chat response.")

heading(doc, "My understanding of Superbrain", 1)
add_para(doc, "I see Superbrain as three layers working together. The IDE is the developer-facing surface where intent, relevant context, and review are visible. The agent turns that intent into a plan, edits, commands, and verification work. The context engine sits underneath: it selects and compresses the repository knowledge the agent needs so the agent can reason across files without loading the full codebase every time.")
add_para(doc, "The key product loop is: a developer describes a change in the IDE, the agent asks the context engine for the relevant architecture and dependencies, the agent proposes or executes work, and the developer reviews the result. ChangeScope focuses on the decision surface between retrieval and execution. It explains why these files matter, what might break, and what proof should be required.")

heading(doc, "Demo workflow", 1)
for text in [
    "Choose one of three realistic changes: add RBAC, add sensitive-action audit events, or move billing webhooks to a worker.",
    "Inspect the dependency path and the selected working context. The product explains what is included and what is intentionally excluded.",
    "Review a four-step implementation plan with file-level rationale and verification criteria.",
    "Approve the plan explicitly. The demo confirms that the plan is ready to hand to an execution agent, while preserving human control."
]:
    p = doc.add_paragraph(style="List Number 2"); p.paragraph_format.space_after = Pt(4); p.add_run(text)

heading(doc, "Architecture and design", 1)
add_para(doc, "ChangeScope is a standalone static application written in vanilla HTML, CSS, and JavaScript. It has no server, database, authentication layer, or external API requirement. This is intentional: for an assignment demo, I wanted the central product interaction to be reliable, inspectable, and deployable at zero marginal cost.")
for label, decision, why in [
    ("Interface", "Single-screen analysis workspace", "Keeps the journey visible: issue, context, plan, proof, and approval."),
    ("Scenario engine", "Three deterministic scenarios", "Lets a reviewer test meaningful paths without API cost, keys, latency, or fabricated model behavior."),
    ("Change map", "Rendered dependency graph", "Makes the implied blast radius concrete and allows the plan to be challenged."),
    ("Trust model", "Evidence and approval are first-class", "Keeps agent output reviewable and controllable for authorization and billing work."),
    ("Deployment", "Static Vercel-compatible build", "Needs no environment variables or backend setup.")
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.1
    set_font(p.add_run(f"{label}. "), size=10.5, color=DARK, bold=True)
    set_font(p.add_run(f"{decision}. "), size=10.5, color=INK, bold=True)
    set_font(p.add_run(why), size=10.5, color=INK)

heading(doc, "Key decisions and tradeoffs", 1)
heading(doc, "1. I optimized for one strong workflow", 2)
add_para(doc, "I did not try to recreate a coding agent. That would create a shallow product with a chat box and little evidence of judgement. Instead, I scoped the product to the decision immediately before execution. The work is deciding the change boundary and defining what proof is needed.")
heading(doc, "2. I did not add an LLM API", 2)
add_para(doc, "The prototype uses deliberately authored demo scenarios rather than a paid model call. This avoids passing off unpredictable output as product reliability. In production, the scenario engine would be replaced by repository indexing, dependency analysis, and an agent planner. The UI contract would stay the same: each conclusion needs a traceable reason, a bounded context, and a verification path.")
heading(doc, "3. I made uncertainty visible", 2)
add_para(doc, "The risk callout, excluded-context explanation, and verification cards are not decoration. They are the parts of the interface that help a developer decide whether to trust a plan or investigate further. A good agent should make a developer faster without making them less informed.")

heading(doc, "If I were building Superbrain next", 1)
add_para(doc, "I would focus next on making the context engine legible to the developer. Context efficiency is valuable, but users need to see why the system selected a set of files and what it may have missed. I would add a lightweight context inspector with: (1) selected and excluded files, (2) dependency-path evidence, (3) confidence and risk signals, and (4) the ability to pin or remove context before execution.")
add_para(doc, "I would also make verification a visible phase of the product. After an agent makes a change, the interface should communicate which tests ran, what passed, what remains unproven, and why the changed files were sufficient. This would make the product feel less like an opaque assistant and more like an engineering collaborator.")

heading(doc, "UI opportunities", 1)
add_para(doc, "My main product preference is for interfaces that make state and control obvious. In this category, a user should always be able to answer: What has the agent understood? What is it about to do? What changed? How can I approve, interrupt, or revise it? A clear plan-review-execute-verify progression would reduce the cognitive load of long agent sessions.")
add_para(doc, "In particular, I would avoid relying only on a linear chat history for complex work. Chat is useful for intent, but it is a poor surface for comparing files, tracking multi-step plans, and understanding system-level impact. Persistent structured panels for context, plan, diffs, and proof would make complex work easier to inspect.")

heading(doc, "How I would extend ChangeScope", 1)
for text in [
    "Replace the scenario data with a repository adapter that indexes imports, routes, tests, ownership, and recent changes.",
    "Add a plan editor so a developer can change assumptions before approval.",
    "Connect an execution agent after approval, with real diffs and a test-run timeline.",
    "Add collaboration: reviewers can comment on a plan, pin required files, and approve high-risk changes."
]:
    p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(4); p.add_run(text)

heading(doc, "How to run", 1)
add_para(doc, "Clone the repository and open index.html directly in a modern browser. For a production bundle, run npm run build. The project is ready for Vercel as a static deployment and requires no environment variables.")

doc.core_properties.title = "Founding AI Engineer Assignment - ChangeScope"
doc.core_properties.author = "Madhavik Dogra"
doc.save(OUT)
print(OUT)
